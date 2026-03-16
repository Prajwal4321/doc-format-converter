import sys
import ollama
import fitz  # PyMuPDF
import pdfplumber
import pandas as pd
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout,
    QHBoxLayout, QPushButton, QTextEdit, QFileDialog,
    QLabel, QSplitter, QMessageBox, QScrollArea, QProgressBar,
    QCheckBox, QGroupBox
)
from PyQt6.QtGui import QPixmap, QImage
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from fpdf import FPDF
from docx import Document

# --- WORKER THREAD FOR AI ---
class AIWorker(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, text):
        super().__init__()
        self.text = text

    def run(self):
        try:
            prompt = (
                f"Analyze this document. Provide a summary of key objectives, "
                f"technical findings, and conclusions:\n\n{self.text[:8000]}"
            )
            response = ollama.generate(model='llama3', prompt=prompt)
            self.finished.emit(response['response'])
        except Exception as e:
            self.error.emit(str(e))

# --- MAIN APPLICATION ---
class OfflineApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Offline AI Document Assistant - Pro Edition")
        self.resize(1400, 900)
        self.file_path = ""
        self.doc = None
        self.current_pdf_text = ""
        self.init_ui()

    def init_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)
        self.splitter = QSplitter(Qt.Orientation.Horizontal)

        # --- LEFT SIDE: Document Viewer ---
        left_container = QWidget()
        left_layout = QVBoxLayout(left_container)
        self.upload_btn = QPushButton(" 📂 Upload a PDF")
        self.upload_btn.setStyleSheet("height: 50px; font-weight: bold; background-color: #000000;")
        self.upload_btn.clicked.connect(self.upload_file)

        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.page_container = QWidget()
        self.page_layout = QVBoxLayout(self.page_container)
        self.scroll_area.setWidget(self.page_container)

        left_layout.addWidget(self.upload_btn)
        left_layout.addWidget(self.scroll_area)

        # --- RIGHT SIDE: Tools & AI ---
        right_container = QWidget()
        right_layout = QVBoxLayout(right_container)

        # AI SECTION (Split into 2 parts)
        ai_section_layout = QHBoxLayout()

        # Part A: Status Bar (Checkboxes)
        self.status_group = QGroupBox("Status Bar")
        status_vbox = QVBoxLayout()
        self.check_upload = QCheckBox("File Uploaded")
        self.check_tables = QCheckBox("Tables Scanned")
        self.check_ai = QCheckBox("AI Analysis Done")
        # Disable manual clicking to make them true status indicators
        for cb in [self.check_upload, self.check_tables, self.check_ai]:
            cb.setEnabled(False)
            status_vbox.addWidget(cb)
        status_vbox.addStretch()
        self.status_group.setLayout(status_vbox)
        self.status_group.setFixedWidth(180)

        # Part B: AI Summary Text
        summary_container = QWidget()
        summary_vbox = QVBoxLayout(summary_container)
        summary_vbox.addWidget(QLabel("🤖 AI Automated Summary:"))
        self.ai_summary = QTextEdit()
        summary_vbox.addWidget(self.ai_summary)

        ai_section_layout.addWidget(self.status_group)
        ai_section_layout.addWidget(summary_container)

        # Progress Indicators
        self.status_label = QLabel("Ready")
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 0)
        self.progress_bar.hide()

        # Tools
        #self.scan_tables_btn = QPushButton(" 📊 Scan for Tables")
        #self.scan_tables_btn.clicked.connect(self.scan_tables)
        self.manual_input = QTextEdit()
        self.manual_input.setPlaceholderText("Manual notes and tables will appear here...")

        # Buttons
        btn_layout = QHBoxLayout()
        self.save_pdf_btn = QPushButton(" 💾 View PDF")
        self.save_docx_btn = QPushButton(" 📝 Save as Word")
        self.save_pdf_btn.clicked.connect(self.export_pdf)
        self.save_docx_btn.clicked.connect(self.export_docx)
        btn_layout.addWidget(self.save_pdf_btn)
        btn_layout.addWidget(self.save_docx_btn)

        # Add all to right layout
        right_layout.addLayout(ai_section_layout)
        right_layout.addWidget(self.status_label)
        right_layout.addWidget(self.progress_bar)
        #right_layout.addWidget(self.scan_tables_btn)
        right_layout.addWidget(QLabel(" ✍️ Manual Data & Tables:"))
        right_layout.addWidget(self.manual_input)
        right_layout.addLayout(btn_layout)

        self.splitter.addWidget(left_container)
        self.splitter.addWidget(right_container)
        layout.addWidget(self.splitter)

    # --- LOGIC ---
    def reset_status(self):
        self.check_upload.setChecked(False)
        self.check_tables.setChecked(False)
        self.check_ai.setChecked(False)

    def upload_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Open Document", "", "Documents (*.pdf *.txt)")
        if path:
            self.file_path = path
            self.reset_status()
            self.clear_preview()
            try:
                if path.lower().endswith('.pdf'):
                    self.doc = fitz.open(path)
                    text = ""
                    for i, page in enumerate(self.doc):
                        text += page.get_text()
                        if i < 10:
                            pix = page.get_pixmap(matrix=fitz.Matrix(1.2, 1.2))
                            img = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format.Format_RGB888).copy()
                            lbl = QLabel(); lbl.setPixmap(QPixmap.fromImage(img).scaledToWidth(600))
                            self.page_layout.addWidget(lbl)
                    self.current_pdf_text = text
                
                self.check_upload.setChecked(True)
                self.start_ai_thread(self.current_pdf_text)
            except Exception as e:
                QMessageBox.critical(self, "Error", str(e))

    def scan_tables(self):
        if not self.file_path or not self.file_path.lower().endswith('.pdf'):
            return
        try:
            table_content = "\n--- EXTRACTED TABLES ---\n"
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        table_content += df.to_string(index=False) + "\n\n"
            
            self.manual_input.append(table_content)
            self.check_tables.setChecked(True)
        except Exception as e:
            print(f"Table Scan Error: {e}")

    def start_ai_thread(self, text):
        self.status_label.setText("AI is analyzing...")
        self.progress_bar.show()
        self.worker = AIWorker(text)
        self.worker.finished.connect(self.handle_ai_done)
        self.worker.start()

    def handle_ai_done(self, summary):
        self.ai_summary.setText(summary)
        self.progress_bar.hide()
        self.status_label.setText("Analysis Complete")
        self.check_ai.setChecked(True)

    def clear_preview(self):
        while self.page_layout.count():
            child = self.page_layout.takeAt(0)
            if child.widget(): child.widget().deleteLater()

    # --- EXPORTS (PDF/DOCX) remain same as previous version ---
    def export_pdf(self):
        save_path, _ = QFileDialog.getSaveFileName(self, "Save PDF", "Report.pdf", "PDF (*.pdf)")
        if save_path:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=11)
            pdf.multi_cell(0, 10, self.ai_summary.toPlainText())
            pdf.output(save_path)

    def export_docx(self):
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Word", "Report.docx", "Word (*.docx)")
        if save_path:
            doc = Document()
            doc.add_heading('Analysis Report', 0)
            doc.add_paragraph(self.ai_summary.toPlainText())
            doc.save(save_path)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = OfflineApp()
    window.show()
    sys.exit(app.exec())
